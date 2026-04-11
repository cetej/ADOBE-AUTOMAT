"""Endpointy pro zapis prekladu zpet do IDML a MAP (Illustrator)."""

import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import logging
from pathlib import Path
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse

from config import EXPORTS_DIR
from models import ProjectPhase, ProjectType
from services.project_store import get_project, save_project
from services.idml_writeback import writeback_idml
from services.map_writeback import writeback_map, preview_map
from services.translation_service import update_translation_memory, write_back_to_termdb

logger = logging.getLogger(__name__)
router = APIRouter(tags=["writeback"])


@router.post("/api/projects/{project_id}/writeback")
async def api_writeback(project_id: str):
    """Zapise preklady zpet do IDML a vrati vysledek."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    if not project.idml_path:
        raise HTTPException(400, "Projekt nema IDML soubor")

    idml_path = Path(project.idml_path)
    if not idml_path.exists():
        raise HTTPException(404, f"IDML soubor nenalezen: {idml_path}")

    # Spocitej kolik elementu ma preklad
    with_czech = [e for e in project.elements if e.czech]
    if not with_czech:
        raise HTTPException(400, "Zadne preklady k zapisu")

    try:
        result = writeback_idml(
            idml_path=idml_path,
            elements=project.elements,
            project_id=project_id,
        )
    except Exception as e:
        logger.error("Writeback selhal: %s", e)
        raise HTTPException(500, f"Writeback selhal: {e}")

    # Auto-save: schválené překlady (status=OK) → translation_memory + terminology DB
    tm_added = update_translation_memory(project.elements)
    termdb_added = write_back_to_termdb(project.elements)
    if tm_added:
        logger.info("Writeback: +%d překladů do translation memory", tm_added)
    if termdb_added:
        logger.info("Writeback: +%d termínů do terminology DB", termdb_added)

    # Ulozit cestu k exportu do projektu
    if result.get("output_path"):
        project.exports["idml_cz"] = result["output_path"]
        project.phase = ProjectPhase.EXPORTED
        save_project(project)

    result["tm_added"] = tm_added
    result["termdb_added"] = termdb_added
    return result


@router.post("/api/projects/{project_id}/writeback/preview")
async def api_writeback_preview(project_id: str):
    """Nahled: kolik textu bude zapsano, kolik chybi."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    total = len(project.elements)
    with_czech = sum(1 for e in project.elements if e.czech)
    with_story = sum(1 for e in project.elements if e.czech and e.story_id)
    no_translation = sum(1 for e in project.elements if not e.czech and e.contents.strip())

    return {
        "total_elements": total,
        "with_translation": with_czech,
        "writable": with_story,
        "missing_translation": no_translation,
        "coverage_pct": round(with_czech / total * 100, 1) if total else 0,
    }


@router.post("/api/projects/{project_id}/writeback-map")
async def api_writeback_map(project_id: str):
    """Zapise preklady zpet do Illustratoru (MAP projekt)."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    if project.type != ProjectType.MAP:
        raise HTTPException(400, "Projekt neni typu MAP")

    with_czech = [e for e in project.elements if e.czech]
    if not with_czech:
        raise HTTPException(400, "Zadne preklady k zapisu")

    try:
        result = writeback_map(elements=project.elements)
    except Exception as e:
        logger.error("MAP writeback selhal: %s", e)
        raise HTTPException(500, f"MAP writeback selhal: {e}")

    # Auto-save: schválené překlady → translation_memory + terminology DB
    tm_added = update_translation_memory(project.elements)
    termdb_added = write_back_to_termdb(project.elements)
    if tm_added:
        logger.info("MAP writeback: +%d překladů do translation memory", tm_added)
    if termdb_added:
        logger.info("MAP writeback: +%d termínů do terminology DB", termdb_added)

    if result.get("changed", 0) > 0:
        project.phase = ProjectPhase.WRITTEN_BACK
        save_project(project)

    result["tm_added"] = tm_added
    result["termdb_added"] = termdb_added
    return result


@router.post("/api/projects/{project_id}/writeback-map/preview")
async def api_writeback_map_preview(project_id: str):
    """Nahled pro MAP writeback — kolik textu bude zapsano."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    return preview_map(project.elements)


# Kategorie pro čistý export (redakční obsah)
_CLEAN_CATEGORIES = {
    "title", "heading", "subtitle", "lead", "body", "main_text",
    "bullet", "caption", "info_boxes", "annotations", "labels",
}


def _clean_paragraphs(project) -> list[dict]:
    """Vrátí spojené odstavce pro čistý export.

    Elementy ze stejného story/layer spojí do jednoho odstavce —
    InDesign dělí text na fragmenty po <Br/> a <Content>,
    ale pro export chceme plynulý text.

    Returns:
        list[dict]: [{"text": "...", "category": "body", "group": "Story_u36e1"}]
    """
    # Filtruj elementy s překladem
    elements = [e for e in project.elements
                if e.czech and e.czech.strip()
                and (not e.category or e.category.value in _CLEAN_CATEGORIES or e.category is None)]

    if not elements:
        return []

    # Seskup fragmenty do odstavců — mezera v indexu = nový odstavec
    # InDesign dělí text na Content elementy (inline formátování),
    # ale <Br/> mezi nimi se projeví jako přeskočený index.
    import re

    paragraphs = []
    current_story = None
    current_texts = []
    current_cats = []
    last_index = -1

    def _flush():
        if not current_texts:
            return
        joined = " ".join(t for t in current_texts if t)
        joined = re.sub(r'\s+', ' ', joined)
        joined = re.sub(r'\s+([.,;:!?)\u201c\u201d])', r'\1', joined)
        cat = ""
        if current_cats:
            from collections import Counter
            cat = Counter(current_cats).most_common(1)[0][0]
        if joined.strip():
            paragraphs.append({"text": joined.strip(), "category": cat, "group": current_story})

    for elem in elements:
        story = elem.story_id or elem.layer_name or "other"
        # Parsuj index z ID (format: "Story_u123/5")
        try:
            idx = int(elem.id.rsplit("/", 1)[1])
        except (ValueError, IndexError):
            idx = -1

        # Nový odstavec pokud: jiný story NEBO mezera v indexu (přeskočený Br)
        if story != current_story or (idx > last_index + 1 and last_index >= 0):
            _flush()
            current_texts = []
            current_cats = []
            current_story = story

        current_texts.append(elem.czech.strip())
        cat = (elem.category.value if elem.category else "") or ""
        if cat:
            current_cats.append(cat)
        last_index = idx

    _flush()

    # Seřadit podle typu sekce
    section_order = ["title", "heading", "subtitle", "lead", "body", "main_text",
                     "bullet", "caption", "info_boxes", "annotations", "labels"]

    def sort_key(p):
        try:
            return section_order.index(p["category"])
        except ValueError:
            return len(section_order)

    return sorted(paragraphs, key=sort_key)


def _section_label(category: str) -> str:
    """Český název sekce pro export."""
    labels = {
        "title": "Titulek", "heading": "Nadpis", "subtitle": "Podtitulek",
        "lead": "Perex", "body": "Text článku", "main_text": "Hlavní text",
        "bullet": "Odrážky", "caption": "Popisky",
        "info_boxes": "Informační boxy", "annotations": "Anotace",
        "labels": "Popisky grafik/map",
    }
    return labels.get(category, category or "Ostatní")


@router.get("/api/projects/{project_id}/export-docx")
async def api_export_docx(project_id: str):
    """Exportuje čistý překlad jako Word dokument — jen redakční text."""
    from docx import Document
    from docx.shared import Pt

    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    paragraphs = _clean_paragraphs(project)
    if not paragraphs:
        raise HTTPException(400, "Žádné přeložené elementy k exportu")

    doc = Document()
    doc.add_heading(project.name, level=1)

    current_section = None
    for para in paragraphs:
        cat = para["category"]
        section = _section_label(cat)
        if section != current_section:
            current_section = section
            doc.add_heading(section, level=2)

        p = doc.add_paragraph()
        run = p.add_run(para["text"])
        if cat in ("title", "heading"):
            run.font.size = Pt(14)
            run.bold = True
        elif cat in ("subtitle", "lead"):
            run.font.size = Pt(12)
            run.italic = True
        elif cat in ("caption", "labels", "annotations"):
            run.font.size = Pt(9)
        else:
            run.font.size = Pt(11)

    export_dir = EXPORTS_DIR / project_id
    export_dir.mkdir(parents=True, exist_ok=True)
    docx_path = export_dir / f"{project_id}_translation.docx"
    doc.save(str(docx_path))

    return FileResponse(
        path=str(docx_path),
        filename=docx_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/api/projects/{project_id}/export-md")
async def api_export_md(project_id: str):
    """Exportuje čistý překlad jako Markdown — jen redakční text."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    paragraphs = _clean_paragraphs(project)
    if not paragraphs:
        raise HTTPException(400, "Žádné přeložené elementy k exportu")

    lines = [f"# {project.name}", ""]

    current_section = None
    for para in paragraphs:
        cat = para["category"]
        section = _section_label(cat)
        if section != current_section:
            current_section = section
            lines.append(f"## {section}")
            lines.append("")

        if cat in ("title", "heading"):
            lines.append(f"### {para['text']}")
        elif cat in ("caption", "labels", "annotations"):
            lines.append(f"*{para['text']}*")
        else:
            lines.append(para["text"])
        lines.append("")

    md_content = "\n".join(lines)

    export_dir = EXPORTS_DIR / project_id
    export_dir.mkdir(parents=True, exist_ok=True)
    md_path = export_dir / f"{project_id}_translation.md"
    md_path.write_text(md_content, encoding="utf-8")

    return FileResponse(
        path=str(md_path),
        filename=md_path.name,
        media_type="text/markdown",
    )


@router.get("/api/projects/{project_id}/download/{export_key}")
async def api_download_export(project_id: str, export_key: str):
    """Stahne exportovany soubor."""
    project = get_project(project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    file_path = project.exports.get(export_key)
    if not file_path:
        raise HTTPException(404, f"Export '{export_key}' nenalezen")

    path = Path(file_path).resolve()
    if not path.is_relative_to(EXPORTS_DIR.resolve()):
        raise HTTPException(403, "Pristup k souboru mimo export adresar")
    if not path.exists():
        raise HTTPException(404, f"Soubor nenalezen: {path}")

    return FileResponse(
        path=str(path),
        filename=path.name,
        media_type="application/octet-stream",
    )
