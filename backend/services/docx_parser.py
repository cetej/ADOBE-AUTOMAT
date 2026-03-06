"""Parser DOCX souboru s českým překladem.

DOCX od překladatele obsahuje:
- Sekvenční české odstavce (vše styl "Normal")
- Stránkové markery: (legenda str. XX), (sloupek str. XX), (citát str. XX), (str. XX)
- Překladatel sloučil fragmenty z IDML zpět do logických vět
"""

import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import re
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

# Regex pro stránkové markery — (legenda str. 44), (sloupek str. 55-56), (citát str. 77) atd.
PAGE_MARKER_RE = re.compile(
    r'^\((?:legenda|sloupek|citát|str\.?)\s+str\.?\s*(\d+)(?:\s*[-–]\s*(\d+))?\)$',
    re.IGNORECASE
)


@dataclass
class DocxSection:
    """Sekce DOCX — skupina odstavců mezi stránkovými markery."""
    section_type: str  # "body", "legenda", "sloupek", "citát"
    page_start: int
    page_end: int  # stejná jako start pokud není rozsah
    paragraphs: list[str] = field(default_factory=list)
    marker_text: str = ""  # původní text markeru

    @property
    def full_text(self) -> str:
        return "\n".join(self.paragraphs)

    @property
    def page_range(self) -> range:
        return range(self.page_start, self.page_end + 1)


@dataclass
class DocxParseResult:
    """Výsledek parsování DOCX."""
    total_paragraphs: int
    sections: list[DocxSection]
    header_paragraphs: list[str]  # odstavce před prvním markerem (obsah, úvod)

    def filter_pages(self, page_min: int, page_max: int) -> list[DocxSection]:
        """Vrátí sekce překrývající zadaný rozsah stránek."""
        result = []
        for s in self.sections:
            if s.page_end >= page_min and s.page_start <= page_max:
                result.append(s)
        return result


def _detect_marker(text: str) -> tuple[str, int, int] | None:
    """Detekuje stránkový marker. Vrací (typ, page_start, page_end) nebo None."""
    text = text.strip()

    # Specifické typy: (legenda str. 44), (sloupek str. 55)
    for marker_type in ["legenda", "sloupek", "citát"]:
        pattern = re.compile(
            rf'^\({marker_type}\s+str\.?\s*(\d+)(?:\s*[-–]\s*(\d+))?\)$',
            re.IGNORECASE
        )
        m = pattern.match(text)
        if m:
            p1 = int(m.group(1))
            p2 = int(m.group(2)) if m.group(2) else p1
            return (marker_type, p1, p2)

    # Generický: (str. 44) nebo (str. 44-45)
    m = re.match(r'^\(str\.?\s*(\d+)(?:\s*[-–]\s*(\d+))?\)$', text, re.IGNORECASE)
    if m:
        p1 = int(m.group(1))
        p2 = int(m.group(2)) if m.group(2) else p1
        return ("body", p1, p2)

    return None


def parse_docx(docx_path: str | Path) -> DocxParseResult:
    """Parsuje DOCX soubor s překladem. Segmentuje podle stránkových markerů."""
    from docx import Document

    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(str(docx_path))
    all_paragraphs = [p.text.strip() for p in doc.paragraphs]
    non_empty = [p for p in all_paragraphs if p]

    logger.info("DOCX: %d celkem odstavců, %d neprázdných", len(all_paragraphs), len(non_empty))

    sections: list[DocxSection] = []
    header_paragraphs: list[str] = []
    current_section: DocxSection | None = None

    for para_text in non_empty:
        marker = _detect_marker(para_text)

        if marker:
            # Nový marker → nová sekce
            section_type, p_start, p_end = marker
            current_section = DocxSection(
                section_type=section_type,
                page_start=p_start,
                page_end=p_end,
                marker_text=para_text,
            )
            sections.append(current_section)
        elif current_section is not None:
            # Odstavec patří do aktuální sekce
            current_section.paragraphs.append(para_text)
        else:
            # Odstavce před prvním markerem (obsah, metadata)
            header_paragraphs.append(para_text)

    logger.info("DOCX: %d sekcí, %d header odstavců", len(sections), len(header_paragraphs))

    return DocxParseResult(
        total_paragraphs=len(non_empty),
        sections=sections,
        header_paragraphs=header_paragraphs,
    )


def get_body_sections(result: DocxParseResult, page_min: int, page_max: int) -> list[DocxSection]:
    """Vrátí body sekce (ne legendy/sloupky) pro daný rozsah stránek.

    Legendy a sloupky se vrací zvlášť — mají jiný matching pattern.
    """
    filtered = result.filter_pages(page_min, page_max)
    return [s for s in filtered if s.section_type == "body"]


def get_all_filtered_sections(result: DocxParseResult, page_min: int, page_max: int) -> list[DocxSection]:
    """Vrátí VŠECHNY sekce (body + legendy + sloupky + citáty) pro daný rozsah stránek."""
    return result.filter_pages(page_min, page_max)
