"""Parsery korektur z různých formátů souborů.

Podporované formáty:
- Excel (.xlsx) — tabulka s 2 sloupci: originál + oprava
- DOCX (.docx) — track changes, tabulka, nebo arrow separátor
- PDF (.pdf) — tabulky nebo textové vzory
"""

import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

import re
import logging
from pathlib import Path

from services.correction_store import CorrectionEntry

logger = logging.getLogger(__name__)


def parse_corrections_file(path: Path, source_type: str) -> list[CorrectionEntry]:
    """Hlavní dispatcher — parsuje soubor podle typu."""
    parsers = {
        "excel": parse_corrections_xlsx,
        "docx": parse_corrections_docx,
        "pdf": parse_corrections_pdf,
    }
    parser = parsers.get(source_type)
    if not parser:
        raise ValueError(f"Nepodporovaný typ: {source_type}")

    entries = parser(path)
    logger.info("Parsed %d korektur z %s (%s)", len(entries), path.name, source_type)
    return entries


# ─── Excel parser ────────────────────────────────────────────

def parse_corrections_xlsx(path: Path) -> list[CorrectionEntry]:
    """Parsuje Excel s korekturami.

    Očekává 2 sloupce: originál (before) a oprava (after).
    Automaticky detekuje header řádek.
    """
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    if not ws:
        return []

    entries = []
    header_keywords = {"originál", "original", "before", "před", "text",
                       "oprava", "corrected", "after", "po", "korektura"}

    skip_first = False
    rows = list(ws.iter_rows(min_row=1, max_col=10, values_only=True))
    if rows:
        first_row = [str(c).strip().lower() for c in rows[0] if c]
        if any(kw in " ".join(first_row) for kw in header_keywords):
            skip_first = True

    for row in rows[1 if skip_first else 0:]:
        cells = [str(c).strip() if c else "" for c in row]
        # Najdi první dva neprázdné sloupce
        non_empty = [(i, v) for i, v in enumerate(cells) if v]
        if len(non_empty) < 2:
            continue

        before = non_empty[0][1]
        after = non_empty[1][1]

        if before == after:
            continue

        entries.append(CorrectionEntry(
            before=before,
            after=after,
            source="excel",
        ))

    wb.close()
    return entries


# ─── DOCX parser ─────────────────────────────────────────────

def parse_corrections_docx(path: Path) -> list[CorrectionEntry]:
    """Parsuje DOCX s korekturami.

    Strategie (v pořadí priority):
    1. Track-changes (strikethrough + nový text)
    2. Tabulka se 2 sloupci
    3. Arrow separátor (→ nebo >)
    """
    import docx

    doc = docx.Document(str(path))

    # Strategie 1: Track changes (strikethrough runs)
    entries = _parse_docx_track_changes(doc)
    if entries:
        logger.info("DOCX: nalezeny track changes (%d oprav)", len(entries))
        return entries

    # Strategie 2: Tabulka
    entries = _parse_docx_tables(doc)
    if entries:
        logger.info("DOCX: nalezena tabulka korektur (%d oprav)", len(entries))
        return entries

    # Strategie 3: Arrow separátor
    entries = _parse_docx_arrows(doc)
    if entries:
        logger.info("DOCX: nalezeny arrow separátory (%d oprav)", len(entries))
        return entries

    logger.warning("DOCX: žádná rozpoznatelná korektura v %s", path.name)
    return []


def _parse_docx_track_changes(doc) -> list[CorrectionEntry]:
    """Extrahuje track changes z DOCX — strikethrough = smazáno, nový text = vloženo."""
    entries = []

    for para in doc.paragraphs:
        deleted_parts = []
        inserted_parts = []
        has_changes = False

        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue

            if run.font.strike:
                deleted_parts.append(text)
                has_changes = True
            else:
                inserted_parts.append(text)

        if has_changes and deleted_parts:
            before = " ".join(deleted_parts)
            after = " ".join(inserted_parts) if inserted_parts else ""
            entries.append(CorrectionEntry(
                before=before,
                after=after,
                source="docx",
            ))

    return entries


def _parse_docx_tables(doc) -> list[CorrectionEntry]:
    """Parsuje tabulky v DOCX — první dva sloupce jako before/after."""
    entries = []

    for table in doc.tables:
        if len(table.columns) < 2:
            continue

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue

            # Přeskoč header
            if i == 0:
                header_test = " ".join(cells).lower()
                if any(kw in header_test for kw in ("originál", "original", "before", "oprava", "after")):
                    continue

            before, after = cells[0], cells[1]
            if not before or before == after:
                continue

            entries.append(CorrectionEntry(
                before=before,
                after=after,
                source="docx",
            ))

    return entries


def _parse_docx_arrows(doc) -> list[CorrectionEntry]:
    """Parsuje odstavce s arrow separátorem: 'starý text → nový text'."""
    entries = []
    arrow_pattern = re.compile(r'(.+?)\s*[→>]\s*(.+)')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        match = arrow_pattern.match(text)
        if match:
            before = match.group(1).strip()
            after = match.group(2).strip()
            if before and after and before != after:
                entries.append(CorrectionEntry(
                    before=before,
                    after=after,
                    source="docx",
                ))

    return entries


# ─── PDF parser ──────────────────────────────────────────────

def parse_corrections_pdf(path: Path) -> list[CorrectionEntry]:
    """Parsuje PDF s korekturami.

    Strategie:
    1. Tabulky (pdfplumber extract_tables)
    2. Textové vzory (arrow separátor)
    """
    import pdfplumber

    entries = []

    with pdfplumber.open(str(path)) as pdf:
        # Strategie 1: Tabulky
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in (tables or []):
                for i, row in enumerate(table):
                    if not row or len(row) < 2:
                        continue
                    cells = [str(c).strip() if c else "" for c in row]

                    # Přeskoč header
                    if i == 0:
                        header_test = " ".join(cells).lower()
                        if any(kw in header_test for kw in ("originál", "original", "before", "oprava", "after")):
                            continue

                    before, after = cells[0], cells[1]
                    if not before or before == after:
                        continue

                    entries.append(CorrectionEntry(
                        before=before,
                        after=after,
                        source="pdf",
                    ))

        if entries:
            logger.info("PDF: nalezeny tabulky korektur (%d oprav)", len(entries))
            return entries

        # Strategie 2: Textové vzory
        arrow_pattern = re.compile(r'(.+?)\s*[→>]\s*(.+)')
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                line = line.strip()
                match = arrow_pattern.match(line)
                if match:
                    before = match.group(1).strip()
                    after = match.group(2).strip()
                    if before and after and before != after:
                        entries.append(CorrectionEntry(
                            before=before,
                            after=after,
                            source="pdf",
                        ))

    if entries:
        logger.info("PDF: nalezeny textové korekturní vzory (%d oprav)", len(entries))

    return entries
